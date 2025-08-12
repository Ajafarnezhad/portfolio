#!/usr/bin/env python3
"""
Ultra-Advanced Word Document Style Analyzer

This script provides a state-of-the-art pipeline for analyzing Microsoft Word (.docx) documents, extracting and categorizing styled text (bold, italic, underlined, and combinations), and offering advanced NLP insights. It supports multi-threaded processing, diverse output formats (append, JSON, CSV, HTML, SQLite), interactive visualizations, and comprehensive reporting. Built with modern Python libraries, optimized for performance, and designed for scalability.

Author: Amirhossein jafarnezhad
Date: August 12, 2025
Version: 2.0.0
License: MIT License[](https://opensource.org/license/mit/)
"""

import argparse
import logging
import sys
import os
import json
import csv
import sqlite3
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from typing import Dict, List, Tuple, Optional
from collections import defaultdict
import unittest
from concurrent.futures import ThreadPoolExecutor
import time
import re
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.sentiment import SentimentIntensityAnalyzer
import plotly.express as px
import plotly.io as pio
pio.renderers.default = 'notebook'  # Ensure Plotly renders in Jupyter
import pandas as pd
from sqlalchemy import create_engine

# Download NLTK data (run once)
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('vader_lexicon')

# Set up structured logging with file and console handlers
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)
file_handler = logging.FileHandler('style_analyzer.log')
file_handler.setLevel(logging.INFO)
file_handler.setFormatter(logging.Formatter("%(asctime)s - %(levelname)s - %(message)s"))
logger.addHandler(file_handler)

class DocumentStyleAnalyzer:
    def __init__(self, input_file: str, output_dir: str = "output", max_workers: int = 8, 
                 output_format: str = "append", db_path: str = "style_analyzer.db"):
        """Initialize the DocumentStyleAnalyzer with advanced configurations."""
        self.input_file = input_file
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)
        self.output_file = os.path.join(self.output_dir, os.path.basename(input_file))
        self.max_workers = max_workers
        self.output_format = output_format
        self.db_path = db_path
        self.document = None
        self.styled_words = defaultdict(list)  # Category: List of words
        self.total_words = 0
        self.unwanted_pattern = re.compile(r'["\'’“:,\n\-\—\.,;!?]')
        self.stop_words = set(stopwords.words('english'))
        self.sentiment_analyzer = SentimentIntensityAnalyzer()
        self.categories = {
            "b": "Bold Words",
            "i": "Italicized Words",
            "u": "Underlined Words",
            "bi": "Bold & Italicized Words",
            "bu": "Bold & Underlined Words",
            "iu": "Italicized & Underlined Words",
            "biu": "Bold & Italicized & Underlined Words",
        }
        self.stats = {"word_count": 0, "styled_count": 0, "sentiment_score": 0.0}

    def load_document(self) -> None:
        """Load the .docx document with error handling and metadata extraction."""
        try:
            self.document = Document(self.input_file)
            logger.info(f"Loaded document: {self.input_file} with {len(self.document.paragraphs)} paragraphs and {sum(len(table.rows) for table in self.document.tables)} table rows")
            self._extract_metadata()
        except Exception as e:
            logger.error(f"Failed to load document: {e}")
            sys.exit(1)

    def _extract_metadata(self) -> None:
        """Extract document metadata (e.g., author, creation date)."""
        core_props = self.document.core_properties
        self.stats["author"] = getattr(core_props, "author", "Unknown")
        self.stats["created"] = getattr(core_props, "created", "Unknown").isoformat() if getattr(core_props, "created", None) else "Unknown"
        logger.info(f"Extracted metadata: Author={self.stats['author']}, Created={self.stats['created']}")

    def analyze_text(self) -> None:
        """Analyze paragraphs and tables for styled text using multi-threading and NLP."""
        if self.document is None:
            self.load_document()
        
        tasks = self.document.paragraphs + [cell for table in self.document.tables for row in table.rows for cell in row.cells]
        
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            list(executor.map(self._process_element, tasks))
        
        self._compute_sentiment()
        logger.info("Text analysis and sentiment computation completed.")

    def _process_element(self, element) -> None:
        """Process a single paragraph or table cell for styled text and NLP features."""
        if hasattr(element, 'runs'):
            for run in element.runs:
                words = word_tokenize(run.text)
                self.total_words += len(words)
                for word in words:
                    if word.lower() in self.stop_words:
                        continue
                    key = ""
                    if run.bold:
                        key += "b"
                    if run.italic:
                        key += "i"
                    if run.underline:
                        key += "u"
                    if key:
                        cleaned_word = self.unwanted_pattern.sub('', word)
                        if cleaned_word and cleaned_word not in self.styled_words[key]:
                            self.styled_words[key].append(cleaned_word)
                self.stats["styled_count"] += sum(1 for run in element.runs if run.bold or run.italic or run.underline)

    def _compute_sentiment(self) -> None:
        """Compute sentiment score for the document using VADER."""
        text = ' '.join(p.text for p in self.document.paragraphs)
        sentiment = self.sentiment_analyzer.polarity_scores(text)
        self.stats["sentiment_score"] = sentiment['compound']
        logger.info(f"Computed sentiment score: {self.stats['sentiment_score']}")

    def write_results(self) -> None:
        """Write the categorized styled words and stats to the output based on format."""
        if self.output_format == "append":
            self._append_to_document()
        elif self.output_format == "json":
            self._export_to_json()
        elif self.output_format == "csv":
            self._export_to_csv()
        elif self.output_format == "html":
            self._export_to_html()
        elif self.output_format == "db":
            self._export_to_database()
        logger.info(f"Results written in {self.output_format} format to {self.output_file}")

    def _append_to_document(self) -> None:
        """Append categorized styled words and stats to the original document with formatting."""
        title_p = self.document.add_paragraph("\n========== Extracted Styled Words & Analysis ==========\n")
        title_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.runs[0].font.size = Pt(16)
        title_p.runs[0].font.color.rgb = RGBColor(0, 0, 255)
        title_p.runs[0].bold = True

        stats_p = self.document.add_paragraph(f"Analysis Summary:\n- Total Words: {self.total_words}\n- Styled Words: {self.stats['styled_count']}\n- Sentiment Score: {self.stats['sentiment_score']:.2f}\n- Author: {self.stats['author']}\n- Created: {self.stats['created']}")
        stats_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        stats_p.runs[0].font.size = Pt(12)

        for key, category in self.categories.items():
            if self.styled_words[key]:
                cat_p = self.document.add_paragraph(f"\n{category}:")
                cat_p.runs[0].font.size = Pt(14)
                cat_p.runs[0].bold = True
                words_p = self.document.add_paragraph(", ".join(sorted(self.styled_words[key])))
                words_p.ragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                words_p.runs[0].font.size = Pt(12)

        ending_p = self.document.add_paragraph("\n===================\n")
        ending_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ending_p.runs[0].font.size = Pt(14)
        ending_p.runs[0].italic = True

        self.document.save(self.output_file)
        logger.info(f"Appended results to {self.output_file}")

    def _export_to_json(self) -> None:
        """Export categorized styled words and stats to JSON."""
        results = {
            "metadata": {k: v for k, v in self.stats.items() if k != "word_count"},
            "total_words": self.total_words,
            "styled_words": {key: sorted(words) for key, words in self.styled_words.items()}
        }
        with open(self.output_file.replace('.docx', '.json'), 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=4, ensure_ascii=False)
        logger.info(f"Exported results to JSON: {self.output_file.replace('.docx', '.json')}")

    def _export_to_csv(self) -> None:
        """Export categorized styled words and stats to CSV."""
        with open(self.output_file.replace('.docx', '.csv'), 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['Category', 'Words'])
            for key, words in self.styled_words.items():
                writer.writerow([self.categories[key], ', '.join(sorted(words))])
            writer.writerow(['Total Words', self.total_words])
            writer.writerow(['Styled Words Count', self.stats['styled_count']])
            writer.writerow(['Sentiment Score', self.stats['sentiment_score']])
            writer.writerow(['Author', self.stats['author']])
            writer.writerow(['Created', self.stats['created']])
        logger.info(f"Exported results to CSV: {self.output_file.replace('.docx', '.csv')}")

    def _export_to_html(self) -> None:
        """Export categorized styled words and stats to HTML."""
        html_content = f"""
        <html>
        <head><title>Style Analysis Report</title><style>body {{font-family: Arial, sans-serif;}} h2 {{color: #1E90FF;}} p {{margin: 10px 0;}}</style></head>
        <body>
        <h2>Style Analysis Report for {os.path.basename(self.input_file)}</h2>
        <p><strong>Total Words:</strong> {self.total_words}</p>
        <p><strong>Styled Words Count:</strong> {self.stats['styled_count']}</p>
        <p><strong>Sentiment Score:</strong> {self.stats['sentiment_score']:.2f}</p>
        <p><strong>Author:</strong> {self.stats['author']}</p>
        <p><strong>Created:</strong> {self.stats['created']}</p>
        """
        for key, category in self.categories.items():
            if self.styled_words[key]:
                html_content += f"<h3>{category}</h3><p>{', '.join(sorted(self.styled_words[key]))}</p>"
        html_content += "</body></html>"
        with open(self.output_file.replace('.docx', '.html'), 'w', encoding='utf-8') as f:
            f.write(html_content)
        logger.info(f"Exported results to HTML: {self.output_file.replace('.docx', '.html')}")

    def _export_to_database(self) -> None:
        """Export categorized styled words and stats to SQLite database."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('''CREATE TABLE IF NOT EXISTS StyleAnalysis 
                          (document_name TEXT, category TEXT, words TEXT, total_words INTEGER, 
                           styled_count INTEGER, sentiment_score REAL, author TEXT, created TEXT)''')
        
        for key, category in self.categories.items():
            words_str = ', '.join(sorted(self.styled_words[key]))
            cursor.execute("INSERT INTO StyleAnalysis VALUES (?, ?, ?, ?, ?, ?, ?, ?)", 
                           (os.path.basename(self.input_file), category, words_str, self.total_words, 
                            self.stats['styled_count'], self.stats['sentiment_score'], self.stats['author'], self.stats['created']))
        conn.commit()
        conn.close()
        logger.info(f"Exported results to SQLite database: {self.db_path}")

    def visualize_statistics(self, output_dir: str = "plots/") -> None:
        """Visualize statistics of styled words with interactive Plotly."""
        os.makedirs(output_dir, exist_ok=True)
        counts = {self.categories[key]: len(words) for key, words in self.styled_words.items()}
        
        # Interactive bar chart
        fig = px.bar(x=list(counts.keys()), y=list(counts.values()), 
                     title="Styled Word Counts by Category", 
                     labels={'x': 'Style Category', 'y': 'Count'}, 
                     color=list(counts.values()), color_continuous_scale='Viridis')
        fig.update_layout(template="plotly_dark", height=500)
        fig.write_html(os.path.join(output_dir, "styled_word_counts.html"))
        fig.show()
        logger.info("Saved interactive styled word counts visualization.")

    def get_analysis_summary(self) -> Dict:
        """Return a summary of the analysis."""
        return {
            "total_words": self.total_words,
            "styled_words_count": self.stats['styled_count'],
            "sentiment_score": self.stats['sentiment_score'],
            "author": self.stats['author'],
            "created": self.stats['created'],
            "styled_words": dict(self.styled_words)
        }

class TestDocumentStyleAnalyzer(unittest.TestCase):
    def setUp(self):
        """Set up test case with a sample document."""
        self.doc = Document()
        p = self.doc.add_paragraph()
        run = p.add_run("Bold text")
        run.bold = True
        run = p.add_run(" Italic text")
        run.italic = True
        run = p.add_run(" Underlined text")
        run.underline = True
        self.test_file = "test.docx"
        self.doc.save(self.test_file)

    def tearDown(self):
        """Clean up test file."""
        if os.path.exists(self.test_file):
            os.remove(self.test_file)

    def test_analysis(self):
        """Test basic analysis functionality."""
        analyzer = DocumentStyleAnalyzer(self.test_file, output_format="json")
        analyzer.load_document()
        analyzer.analyze_text()
        self.assertIn("b", analyzer.styled_words)
        self.assertIn("Bold", analyzer.styled_words["b"])
        self.assertEqual(analyzer.get_total_words(), 3)

    def test_sentiment(self):
        """Test sentiment analysis functionality."""
        analyzer = DocumentStyleAnalyzer(self.test_file)
        analyzer.load_document()
        analyzer.analyze_text()
        self.assertIsInstance(analyzer.stats["sentiment_score"], float)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Ultra-Advanced Word Document Style Analyzer.")
    parser.add_argument("input_file", type=str, help="Path to the input .docx file.")
    parser.add_argument("--output_dir", type=str, default="output", help="Directory for output files (default: output).")
    parser.add_argument("--output_format", type=str, choices=["append", "json", "csv", "html", "db"], default="append", 
                        help="Output format: append to docx, json, csv, html, or database.")
    parser.add_argument("--max_workers", type=int, default=8, help="Number of concurrent workers for analysis.")
    parser.add_argument("--visualize", action="store_true", help="Generate interactive visualizations.")
    parser.add_argument("--run_tests", action="store_true", help="Run unit tests.")
    parser.add_argument("--verbose", action="store_true", help="Enable verbose output.")
    args = parser.parse_args()

    if args.verbose:
        logger.setLevel(logging.DEBUG)

    analyzer = DocumentStyleAnalyzer(args.input_file, args.output_dir, args.max_workers, args.output_format)
    start_time = time.time()
    analyzer.load_document()
    analyzer.analyze_text()
    analyzer.write_results()
    
    if args.visualize:
        analyzer.visualize_statistics()

    summary = analyzer.get_analysis_summary()
    logger.info(f"Analysis Summary: {summary}")
    logger.info(f"Execution time: {time.time() - start_time:.2f} seconds")

    if args.run_tests:
        unittest.main(argv=[sys.argv[0]])